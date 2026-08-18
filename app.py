import streamlit as st
from io import BytesIO
from types import SimpleNamespace
import re
import sqlite3
from datetime import datetime
import hashlib
import secrets

import plotly.graph_objects as go
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from xml.sax.saxutils import escape

from PyPDF2 import PdfReader
from docx import Document

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from sentence_transformers import SentenceTransformer


# --------------------------------------------------
# PAGE CONFIGURATION
# --------------------------------------------------

st.set_page_config(
    page_title="Plagiarism Detection System",
    page_icon="🔍",
    layout="wide"
)


# --------------------------------------------------
# SESSION STATE
# --------------------------------------------------

DEFAULTS = {
    "uploaded_name1": "",
    "uploaded_name2": "",
    "uploaded_text1": "",
    "uploaded_text2": "",
    "logged_in": False,
    "user_id": None,
    "username": "",
}

for key, value in DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = value


# --------------------------------------------------
# SQLITE DATABASE + USER AUTHENTICATION
# --------------------------------------------------

DB_NAME = "plagiarism_history.db"


def hash_password(password, salt=None):
    if salt is None:
        salt = secrets.token_hex(16)

    password_hash = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        120000
    ).hex()

    return salt, password_hash


def verify_password(password, salt, stored_hash):
    _, password_hash = hash_password(password, salt)
    return secrets.compare_digest(
        password_hash,
        stored_hash
    )


def init_database():
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            salt TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS analysis_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            original_file TEXT NOT NULL,
            checked_file TEXT NOT NULL,
            similarity REAL NOT NULL,
            risk TEXT NOT NULL,
            similar_sentences INTEGER NOT NULL,
            created_at TEXT NOT NULL
        )
    """)

    # Upgrade existing Step-17/18/19 databases.
    columns = [
        row[1]
        for row in cursor.execute(
            "PRAGMA table_info(analysis_history)"
        ).fetchall()
    ]

    if "user_id" not in columns:
        cursor.execute(
            "ALTER TABLE analysis_history ADD COLUMN user_id INTEGER"
        )

    conn.commit()
    conn.close()


def register_user(username, password):
    username = username.strip()

    if not username or not password:
        return False, "Username and password are required."

    if len(username) < 3:
        return False, "Username must contain at least 3 characters."

    if len(password) < 6:
        return False, "Password must contain at least 6 characters."

    salt, password_hash = hash_password(password)

    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO users
            (username, password_hash, salt, created_at)
            VALUES (?, ?, ?, ?)
        """, (
            username,
            password_hash,
            salt,
            datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        ))

        conn.commit()
        user_id = cursor.lastrowid

        # Preserve existing history when the first account is created.
        total_users = cursor.execute(
            "SELECT COUNT(*) FROM users"
        ).fetchone()[0]

        if total_users == 1:
            cursor.execute(
                "UPDATE analysis_history SET user_id = ? "
                "WHERE user_id IS NULL",
                (user_id,)
            )
            conn.commit()

        conn.close()

        return True, "Account created successfully."

    except sqlite3.IntegrityError:
        return False, "Username already exists."


def authenticate_user(username, password):
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    row = cursor.execute("""
        SELECT *
        FROM users
        WHERE username = ?
    """, (username.strip(),)).fetchone()

    conn.close()

    if row is None:
        return None

    if verify_password(
        password,
        row["salt"],
        row["password_hash"]
    ):
        return row

    return None


def save_analysis(
    original_file,
    checked_file,
    similarity,
    risk,
    similar_count,
    user_id
):
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()

    cursor.execute("""
        INSERT INTO analysis_history
        (original_file, checked_file, similarity, risk,
         similar_sentences, created_at, user_id)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        original_file,
        checked_file,
        similarity,
        risk,
        similar_count,
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        user_id
    ))

    conn.commit()
    conn.close()


def get_history(user_id):
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    rows = cursor.execute("""
        SELECT *
        FROM analysis_history
        WHERE user_id = ?
        ORDER BY id DESC
    """, (user_id,)).fetchall()

    conn.close()
    return rows


def delete_analysis(record_id, user_id):
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()

    cursor.execute("""
        DELETE FROM analysis_history
        WHERE id = ? AND user_id = ?
    """, (record_id, user_id))

    conn.commit()
    conn.close()


def clear_history(user_id):
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()

    cursor.execute("""
        DELETE FROM analysis_history
        WHERE user_id = ?
    """, (user_id,))

    conn.commit()
    conn.close()


init_database()


# --------------------------------------------------
# PROFESSIONAL UI STYLING
# --------------------------------------------------

st.markdown("""
<style>

.stApp {
    background-color: #f7f9fc;
}

.main-title {
    text-align: center;
    font-size: 42px;
    font-weight: 700;
    margin-bottom: 5px;
}

.main-subtitle {
    text-align: center;
    color: #64748b;
    font-size: 17px;
    margin-bottom: 35px;
}

h2, h3 {
    color: #172033;
}

[data-testid="stFileUploader"] {
    background-color: white;
    border: 2px dashed #cbd5e1;
    border-radius: 14px;
    padding: 15px;
}

[data-testid="stFileUploader"]:hover {
    border-color: #6366f1;
}

[data-testid="stMetric"] {
    background-color: white;
    border-radius: 14px;
    padding: 18px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 3px 10px rgba(0, 0, 0, 0.05);
}

.stButton > button {
    width: 100%;
    border-radius: 10px;
    font-weight: 600;
    padding: 12px;
    border: none;
    background-color: #4f46e5;
    color: white;
}

.stButton > button:hover {
    background-color: #4338ca;
    color: white;
}

.stDownloadButton > button {
    width: 100%;
    border-radius: 10px;
    font-weight: 600;
    padding: 12px;
}

[data-testid="stExpander"] {
    background-color: white;
    border-radius: 12px;
    border: 1px solid #e2e8f0;
}

[data-testid="stAlert"] {
    border-radius: 10px;
}

hr {
    margin-top: 30px;
    margin-bottom: 30px;
}

.footer {
    text-align: center;
    color: #64748b;
    font-size: 14px;
    margin-top: 40px;
    padding: 20px;
}

</style>
""", unsafe_allow_html=True)


# --------------------------------------------------
# LOGIN / REGISTER
# --------------------------------------------------

if not st.session_state.logged_in:

    st.markdown(
        '<div class="main-title">🔐 Plagiarism Checker</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="main-subtitle">'
        'Secure login for your plagiarism analysis workspace'
        '</div>',
        unsafe_allow_html=True
    )

    login_tab, register_tab = st.tabs(
        ["🔑 Login", "📝 Create Account"]
    )

    with login_tab:

        with st.form("login_form"):
            login_username = st.text_input(
                "Username"
            )

            login_password = st.text_input(
                "Password",
                type="password"
            )

            login_submit = st.form_submit_button(
                "🔐 Login",
                use_container_width=True
            )

        if login_submit:

            user = authenticate_user(
                login_username,
                login_password
            )

            if user:

                st.session_state.logged_in = True
                st.session_state.user_id = user["id"]
                st.session_state.username = user["username"]

                st.success(
                    "✅ Login successful."
                )

                st.rerun()

            else:

                st.error(
                    "❌ Invalid username or password."
                )

    with register_tab:

        with st.form("register_form"):
            new_username = st.text_input(
                "Choose username"
            )

            new_password = st.text_input(
                "Choose password",
                type="password"
            )

            confirm_password = st.text_input(
                "Confirm password",
                type="password"
            )

            register_submit = st.form_submit_button(
                "📝 Create Account",
                use_container_width=True
            )

        if register_submit:

            if new_password != confirm_password:

                st.error(
                    "❌ Passwords do not match."
                )

            else:

                success, message = register_user(
                    new_username,
                    new_password
                )

                if success:

                    st.success(
                        "✅ Account created. "
                        "You can now login."
                    )

                else:

                    st.error(
                        f"❌ {message}"
                    )

    st.info(
        "Your password is stored as a salted PBKDF2-SHA256 hash, "
        "not as plain text."
    )

    st.stop()


# --------------------------------------------------
# SIDEBAR NAVIGATION
# --------------------------------------------------

st.sidebar.title("🔍 Plagiarism Checker")

st.sidebar.success(
    f"👤 Logged in as **{st.session_state.username}**"
)

if st.sidebar.button(
    "🚪 Logout",
    use_container_width=True
):
    st.session_state.logged_in = False
    st.session_state.user_id = None
    st.session_state.username = ""

    # Clear uploaded documents on logout.
    st.session_state.uploaded_name1 = ""
    st.session_state.uploaded_name2 = ""
    st.session_state.uploaded_text1 = ""
    st.session_state.uploaded_text2 = ""

    st.rerun()


page = st.sidebar.radio(
    "Navigate",
    [
        "🏠 Home",
        "📄 Upload Documents",
        "📊 Analysis",
        "📜 History"
    ]
)

st.sidebar.divider()

st.sidebar.info(
    "AI-powered document similarity "
    "and plagiarism analysis."
)


# --------------------------------------------------
# TEXT EXTRACTION
# --------------------------------------------------

def extract_text(uploaded_file):
    uploaded_file.seek(0)

    file_type = uploaded_file.name.split(".")[-1].lower()

    if file_type == "txt":
        try:
            return uploaded_file.read().decode("utf-8")
        except UnicodeDecodeError:
            uploaded_file.seek(0)
            return uploaded_file.read().decode("latin-1")

    if file_type == "pdf":
        pdf = PdfReader(uploaded_file)
        text = ""

        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        return text

    if file_type == "docx":
        document = Document(uploaded_file)
        return "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )

    return ""


# --------------------------------------------------
# SIMILAR SENTENCE DETECTION
# --------------------------------------------------

def find_similar_sentences(text1, text2, threshold=0.50):
    sentences1 = re.split(r'(?<=[.!?])\s+', text1.strip())
    sentences2 = re.split(r'(?<=[.!?])\s+', text2.strip())

    sentences1 = [
        sentence.strip()
        for sentence in sentences1
        if sentence.strip()
    ]

    sentences2 = [
        sentence.strip()
        for sentence in sentences2
        if sentence.strip()
    ]

    if not sentences1 or not sentences2:
        return []

    vectorizer = TfidfVectorizer(
        lowercase=True,
        stop_words="english"
    )

    try:
        vectors = vectorizer.fit_transform(
            sentences1 + sentences2
        )
    except ValueError:
        return []

    vectors1 = vectors[:len(sentences1)]
    vectors2 = vectors[len(sentences1):]

    matrix = cosine_similarity(vectors1, vectors2)

    matches = []

    for i, sentence1 in enumerate(sentences1):
        best_index = matrix[i].argmax()
        score = matrix[i][best_index]

        if score >= threshold:
            matches.append({
                "original": sentence1,
                "matched": sentences2[best_index],
                "score": score * 100
            })

    return matches



# --------------------------------------------------
# SEMANTIC SIMILARITY (NLP EMBEDDINGS)
# --------------------------------------------------

@st.cache_resource
def load_semantic_model():
    return SentenceTransformer("all-MiniLM-L6-v2")


def split_sentences(text):
    sentences = re.split(
        r'(?<=[.!?])\s+',
        text.strip()
    )

    return [
        sentence.strip()
        for sentence in sentences
        if sentence.strip()
    ]


def calculate_semantic_similarity(text1, text2):
    sentences1 = split_sentences(text1)
    sentences2 = split_sentences(text2)

    if not sentences1 or not sentences2:
        return 0.0, []

    model = load_semantic_model()

    embeddings1 = model.encode(
        sentences1,
        convert_to_numpy=True,
        normalize_embeddings=True
    )

    embeddings2 = model.encode(
        sentences2,
        convert_to_numpy=True,
        normalize_embeddings=True
    )

    matrix = cosine_similarity(
        embeddings1,
        embeddings2
    )

    best_matches = []

    for i, sentence1 in enumerate(sentences1):
        best_index = matrix[i].argmax()
        semantic_score = float(matrix[i][best_index])

        best_matches.append({
            "original": sentence1,
            "matched": sentences2[best_index],
            "semantic_score": semantic_score * 100
        })

    overall_semantic_score = float(matrix.max()) * 100

    return overall_semantic_score, best_matches


def find_semantic_matches(text1, text2, threshold=0.70):
    _, matches = calculate_semantic_similarity(
        text1,
        text2
    )

    return [
        {
            "original": item["original"],
            "matched": item["matched"],
            "score": item["semantic_score"]
        }
        for item in matches
        if item["semantic_score"] >= threshold * 100
    ]


# --------------------------------------------------
# DOCX REPORT
# --------------------------------------------------

def create_docx_report(
    original_name,
    checked_name,
    similarity_percentage,
    status,
    similar_sentences
):
    document = Document()
    document.add_heading("PLAGIARISM DETECTION REPORT", level=0)
    document.add_paragraph("Generated by Plagiarism Detection System")

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    details = [
        ("Original Document", original_name),
        ("Document Checked", checked_name),
        ("Similarity Score", f"{similarity_percentage:.2f}%"),
        ("Plagiarism Risk", status),
        ("Similar Sentences", str(len(similar_sentences))),
    ]

    for label, value in details:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    document.add_paragraph()
    document.add_heading("Similar Sentences", level=1)

    if similar_sentences:
        for index, item in enumerate(similar_sentences, start=1):
            document.add_heading(
                f"Match {index} — {item['score']:.2f}% similarity",
                level=2
            )
            document.add_paragraph(
                "Original Sentence: " + item["original"]
            )
            document.add_paragraph(
                "Matched Sentence: " + item["matched"]
            )
    else:
        document.add_paragraph("No highly similar sentences were found.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --------------------------------------------------
# PDF REPORT
# --------------------------------------------------

def create_pdf_report(
    original_name,
    checked_name,
    similarity_percentage,
    status,
    similar_sentences
):
    buffer = BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()
    elements = []

    elements.append(
        Paragraph(
            "PLAGIARISM DETECTION REPORT",
            styles["Title"]
        )
    )

    elements.append(Spacer(1, 20))

    document_data = [
        ["Original Document", escape(original_name)],
        ["Document Checked", escape(checked_name)],
        ["Similarity Score", f"{similarity_percentage:.2f}%"],
        ["Plagiarism Risk", status],
        ["Similar Sentences", str(len(similar_sentences))]
    ]

    table = Table(document_data, colWidths=[170, 300])

    table.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("PADDING", (0, 0), (-1, -1), 8)
        ])
    )

    elements.append(table)
    elements.append(Spacer(1, 25))

    elements.append(
        Paragraph(
            "Similar Sentences",
            styles["Heading2"]
        )
    )

    if similar_sentences:
        for index, item in enumerate(similar_sentences, start=1):
            elements.append(
                Paragraph(
                    f"<b>Match {index}</b> - "
                    f"{item['score']:.2f}% similarity",
                    styles["BodyText"]
                )
            )

            elements.append(Spacer(1, 8))

            elements.append(
                Paragraph(
                    "<b>Original Sentence:</b>",
                    styles["BodyText"]
                )
            )

            elements.append(
                Paragraph(
                    escape(item["original"]),
                    styles["BodyText"]
                )
            )

            elements.append(Spacer(1, 8))

            elements.append(
                Paragraph(
                    "<b>Matched Sentence:</b>",
                    styles["BodyText"]
                )
            )

            elements.append(
                Paragraph(
                    escape(item["matched"]),
                    styles["BodyText"]
                )
            )

            elements.append(Spacer(1, 18))
    else:
        elements.append(
            Paragraph(
                "No highly similar sentences were found.",
                styles["BodyText"]
            )
        )

    elements.append(Spacer(1, 20))

    elements.append(
        Paragraph(
            "Generated by: Plagiarism Detection System",
            styles["BodyText"]
        )
    )

    document.build(elements)

    buffer.seek(0)
    return buffer


# --------------------------------------------------
# HISTORY CSV EXPORT
# --------------------------------------------------

def history_to_csv(history):
    lines = [
        "ID,Original Document,Checked Document,Similarity,Risk,Similar Sentences,Created At"
    ]

    for row in history:
        values = [
            row["id"],
            row["original_file"],
            row["checked_file"],
            f"{row['similarity']:.2f}%",
            row["risk"],
            row["similar_sentences"],
            row["created_at"],
        ]

        escaped = []
        for value in values:
            value = str(value).replace('"', '""')
            escaped.append(f'"{value}"')

        lines.append(",".join(escaped))

    return "\n".join(lines)


# --------------------------------------------------
# HOME PAGE
# --------------------------------------------------

if page == "🏠 Home":

    st.markdown(
        '<div class="main-title">'
        '🔍 Plagiarism Detection System'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="main-subtitle">'
        'AI-powered document similarity and plagiarism analysis'
        '</div>',
        unsafe_allow_html=True
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        st.info(
            "📄\n\n"
            "**Upload Documents**\n\n"
            "Upload TXT, PDF or DOCX files."
        )

    with col2:
        st.info(
            "🔍\n\n"
            "**Detect Similarity**\n\n"
            "Compare documents using NLP."
        )

    with col3:
        st.info(
            "📊\n\n"
            "**View Results**\n\n"
            "Get similarity score and risk level."
        )

    st.divider()

    st.subheader("🔐 Secure User Workspace")

    st.write(
        "Each account has its own plagiarism analysis history. "
        "Your password is protected using salted PBKDF2-SHA256 hashing."
    )

    st.subheader("🧠 AI/NLP Technology")

    st.write(
        "The system combines traditional TF-IDF lexical similarity "
        "with Sentence Transformer embeddings to detect both "
        "word-level and meaning-level similarity."
    )

    st.subheader("🚀 How It Works")

    steps = [
        "Upload an original document.",
        "Upload the document you want to check.",
        "Run plagiarism analysis.",
        "Review similarity and matching sentences.",
        "Download the plagiarism report."
    ]

    for number, step in enumerate(steps, start=1):
        st.write(f"**{number}.** {step}")


# --------------------------------------------------
# UPLOAD PAGE
# --------------------------------------------------

elif page == "📄 Upload Documents":

    st.subheader("📄 Upload Documents")

    col1, col2 = st.columns(2)

    with col1:
        st.write("### Original Document")

        file1 = st.file_uploader(
            "Upload original file",
            type=["txt", "pdf", "docx"],
            key="file1"
        )

        if file1 is not None:
            extracted1 = extract_text(file1)

            if extracted1.strip():
                st.session_state.uploaded_name1 = file1.name
                st.session_state.uploaded_text1 = extracted1

    with col2:
        st.write("### Document to Check")

        file2 = st.file_uploader(
            "Upload file to check",
            type=["txt", "pdf", "docx"],
            key="file2"
        )

        if file2 is not None:
            extracted2 = extract_text(file2)

            if extracted2.strip():
                st.session_state.uploaded_name2 = file2.name
                st.session_state.uploaded_text2 = extracted2

    if st.session_state.uploaded_name1:
        st.success(
            f"✅ Original document ready: "
            f"{st.session_state.uploaded_name1}"
        )

        with st.expander("📖 View Original Extracted Text"):
            st.text_area(
                "Original Text",
                st.session_state.uploaded_text1,
                height=250,
                key="saved_original_text"
            )

    if st.session_state.uploaded_name2:
        st.success(
            f"✅ Document to check ready: "
            f"{st.session_state.uploaded_name2}"
        )

        with st.expander("📖 View Text to Check"):
            st.text_area(
                "Text to Check",
                st.session_state.uploaded_text2,
                height=250,
                key="saved_checked_text"
            )

    if (
        st.session_state.uploaded_name1
        and st.session_state.uploaded_name2
    ):
        st.success(
            "✅ Both documents are saved. "
            "Now open 📊 Analysis from the sidebar."
        )

        if st.button(
            "🗑️ Clear Uploaded Documents",
            use_container_width=True
        ):
            st.session_state.uploaded_name1 = ""
            st.session_state.uploaded_name2 = ""
            st.session_state.uploaded_text1 = ""
            st.session_state.uploaded_text2 = ""
            st.rerun()


# --------------------------------------------------
# ANALYSIS PAGE
# --------------------------------------------------

elif page == "📊 Analysis":

    original_name = st.session_state.get(
        "uploaded_name1",
        ""
    )

    checked_name = st.session_state.get(
        "uploaded_name2",
        ""
    )

    text1 = st.session_state.get(
        "uploaded_text1",
        ""
    )

    text2 = st.session_state.get(
        "uploaded_text2",
        ""
    )

    st.subheader("📊 Plagiarism Analysis")

    if not original_name or not checked_name:
        st.warning(
            "⚠️ Please upload both documents first "
            "from the '📄 Upload Documents' page."
        )
        st.info(
            "Use the sidebar to open Upload Documents."
        )

    else:

        st.caption(
            f"📄 Original: {original_name}  |  "
            f"📄 Checked: {checked_name}"
        )

        st.divider()

        if st.button(
            "🔍 Check Plagiarism",
            use_container_width=True
        ):

            if not text1.strip() or not text2.strip():
                st.error(
                    "❌ Could not extract text from "
                    "one or both files."
                )

            else:

                # --------------------------------------------------
                # TF-IDF SIMILARITY
                # --------------------------------------------------

                vectorizer = TfidfVectorizer(
                    lowercase=True,
                    stop_words="english"
                )

                try:
                    vectors = vectorizer.fit_transform(
                        [text1, text2]
                    )

                    tfidf_similarity = cosine_similarity(
                        vectors[0],
                        vectors[1]
                    )[0][0] * 100

                except ValueError:
                    tfidf_similarity = 0.0

                # --------------------------------------------------
                # SEMANTIC SIMILARITY
                # --------------------------------------------------

                try:
                    semantic_similarity, semantic_matches = (
                        calculate_semantic_similarity(
                            text1,
                            text2
                        )
                    )

                    # Combined score:
                    # 40% lexical TF-IDF + 60% semantic meaning
                    combined_similarity = (
                        (tfidf_similarity * 0.40)
                        + (semantic_similarity * 0.60)
                    )

                    similar_sentences = [
                        {
                            "original": item["original"],
                            "matched": item["matched"],
                            "score": item["semantic_score"]
                        }
                        for item in semantic_matches
                        if item["semantic_score"] >= 70
                    ]

                    similarity_method = (
                        "TF-IDF + Semantic Embeddings"
                    )

                except Exception as semantic_error:
                    # Safe fallback if the embedding model cannot load.
                    semantic_similarity = 0.0
                    combined_similarity = tfidf_similarity

                    similar_sentences = find_similar_sentences(
                        text1,
                        text2,
                        threshold=0.50
                    )

                    similarity_method = (
                        "TF-IDF (Semantic model unavailable)"
                    )

                    st.warning(
                        "⚠️ Semantic model could not be loaded. "
                        "TF-IDF analysis was used instead."
                    )

                similarity_percentage = combined_similarity

                if similarity_percentage >= 70:
                    status = "High"
                    risk_message = (
                        "🔴 High Similarity - "
                        "Possible Plagiarism"
                    )
                elif similarity_percentage >= 40:
                    status = "Moderate"
                    risk_message = (
                        "🟡 Moderate Similarity"
                    )
                else:
                    status = "Low"
                    risk_message = (
                        "🟢 Low Similarity - "
                        "Likely Original"
                    )

                # Save permanently to SQLite
                save_analysis(
                    original_name,
                    checked_name,
                    similarity_percentage,
                    status,
                    len(similar_sentences),
                    st.session_state.user_id
                )

                st.success(
                    "✅ Analysis completed and saved to SQLite."
                )

                st.subheader(
                    "📊 Plagiarism Detection Result"
                )

                st.caption(
                    f"Analysis Method: **{similarity_method}**"
                )

                score_col1, score_col2, score_col3 = st.columns(3)

                with score_col1:
                    st.metric(
                        "TF-IDF Score",
                        f"{tfidf_similarity:.2f}%"
                    )

                with score_col2:
                    st.metric(
                        "Semantic Score",
                        f"{semantic_similarity:.2f}%"
                    )

                with score_col3:
                    st.metric(
                        "Combined Score",
                        f"{similarity_percentage:.2f}%"
                    )

                result_col1, result_col2, result_col3 = st.columns(3)

                with result_col1:
                    st.metric(
                        "Similarity Score",
                        f"{similarity_percentage:.2f}%"
                    )

                with result_col2:
                    st.metric(
                        "Similar Sentences",
                        len(similar_sentences)
                    )

                with result_col3:
                    st.metric(
                        "Plagiarism Risk",
                        status
                    )

                st.write("### Risk Assessment")

                if status == "High":
                    st.error(risk_message)
                elif status == "Moderate":
                    st.warning(risk_message)
                else:
                    st.success(risk_message)

                st.write("### 📊 Similarity Level")

                st.progress(
                    min(int(similarity_percentage), 100)
                )

                st.caption(
                    f"Similarity: "
                    f"{similarity_percentage:.2f}%"
                )

                st.divider()

                st.subheader(
                    "📈 Visual Similarity Dashboard"
                )

                fig = go.Figure(
                    go.Indicator(
                        mode="gauge+number",
                        value=similarity_percentage,
                        number={"suffix": "%"},
                        title={
                            "text": "Overall Similarity"
                        },
                        gauge={
                            "axis": {"range": [0, 100]},
                            "steps": [
                                {"range": [0, 40]},
                                {"range": [40, 70]},
                                {"range": [70, 100]}
                            ],
                            "threshold": {
                                "line": {"width": 4},
                                "thickness": 0.75,
                                "value": similarity_percentage
                            }
                        }
                    )
                )

                fig.update_layout(
                    height=400,
                    margin=dict(
                        l=30,
                        r=30,
                        t=80,
                        b=30
                    )
                )

                st.plotly_chart(
                    fig,
                    use_container_width=True
                )

                st.divider()

                st.subheader(
                    "🔎 Similar Sentences"
                )

                if similar_sentences:

                    st.write(
                        f"Found **{len(similar_sentences)}** "
                        "potentially similar sentence(s)."
                    )

                    for index, item in enumerate(
                        similar_sentences,
                        start=1
                    ):

                        with st.expander(
                            f"Match {index} - "
                            f"{item['score']:.2f}% similarity"
                        ):

                            st.write(
                                "**Original Sentence:**"
                            )

                            st.info(item["original"])

                            st.write(
                                "**Matched Sentence:**"
                            )

                            st.warning(item["matched"])

                            st.write(
                                "Sentence Similarity: "
                                f"**{item['score']:.2f}%**"
                            )

                else:
                    st.success(
                        "✅ No highly similar sentences found."
                    )

                st.divider()

                st.subheader(
                    "📋 Analysis Summary"
                )

                summary_col1, summary_col2, summary_col3 = st.columns(3)

                with summary_col1:
                    st.metric(
                        "Overall Similarity",
                        f"{similarity_percentage:.2f}%"
                    )

                with summary_col2:
                    st.metric(
                        "Similar Sentences",
                        len(similar_sentences)
                    )

                with summary_col3:
                    st.metric(
                        "Plagiarism Risk",
                        status
                    )

                st.divider()

                st.divider()

                # --------------------------------------------------
                # ADVANCED PLAGIARISM REPORT
                # --------------------------------------------------

                st.subheader("📋 Advanced Plagiarism Report")

                report_col1, report_col2 = st.columns(2)

                with report_col1:
                    st.write("### 📄 Documents")
                    st.write(
                        f"**Original:** {original_name}"
                    )
                    st.write(
                        f"**Checked:** {checked_name}"
                    )

                with report_col2:
                    st.write("### 🤖 Analysis Method")
                    st.write(
                        "**TF-IDF + Sentence Transformer Embeddings**"
                    )
                    st.write(
                        "**Weighting:** 40% TF-IDF + 60% Semantic"
                    )

                st.divider()

                st.write("### 📊 Score Breakdown")

                breakdown_col1, breakdown_col2, breakdown_col3 = st.columns(3)

                with breakdown_col1:
                    st.metric(
                        "TF-IDF Score",
                        f"{tfidf_similarity:.2f}%"
                    )

                with breakdown_col2:
                    st.metric(
                        "Semantic Score",
                        f"{semantic_similarity:.2f}%"
                    )

                with breakdown_col3:
                    st.metric(
                        "Final Combined Score",
                        f"{similarity_percentage:.2f}%"
                    )

                st.write("### ⚠️ Risk Interpretation")

                if status == "High":
                    st.error(
                        "High plagiarism risk: the documents show "
                        "strong lexical and/or semantic similarity."
                    )
                elif status == "Moderate":
                    st.warning(
                        "Moderate plagiarism risk: meaningful "
                        "similarity was detected and should be reviewed."
                    )
                else:
                    st.success(
                        "Low plagiarism risk: the detected similarity "
                        "is relatively low."
                    )

                st.write("### 🔎 Matching Evidence")

                if similar_sentences:
                    for index, item in enumerate(
                        similar_sentences,
                        start=1
                    ):
                        with st.container(border=True):
                            st.write(
                                f"**Match {index} — "
                                f"{item['score']:.2f}% semantic similarity**"
                            )

                            evidence_col1, evidence_col2 = st.columns(2)

                            with evidence_col1:
                                st.write("**Original Sentence**")
                                st.info(item["original"])

                            with evidence_col2:
                                st.write("**Matched Sentence**")
                                st.warning(item["matched"])

                            st.caption(
                                "This match was selected because its "
                                "semantic similarity exceeded the "
                                "70% review threshold."
                            )
                else:
                    st.success(
                        "No sentence-level semantic matches above "
                        "the 70% review threshold were found."
                    )

                st.write("### 🧾 Final Assessment")

                st.markdown(
                    f"""
                    **Overall Similarity:** {similarity_percentage:.2f}%  
                    **Plagiarism Risk:** {status}  
                    **Similar Sentences:** {len(similar_sentences)}  
                    **Analysis Method:** TF-IDF + Semantic Embeddings
                    """
                )

                st.subheader(
                    "📥 Export Reports"
                )

                report = f"""
PLAGIARISM DETECTION REPORT
===========================

Original Document:
{original_name}

Document Checked:
{checked_name}

-----------------------------------

OVERALL ANALYSIS

TF-IDF Score:
{tfidf_similarity:.2f}%

Semantic Score:
{semantic_similarity:.2f}%

Combined Similarity Score:
{similarity_percentage:.2f}%

Plagiarism Risk:
{status}

Similar Sentences:
{len(similar_sentences)}

Analysis Method:
TF-IDF + Sentence Transformer Embeddings

Weighting:
40% TF-IDF + 60% Semantic Similarity

-----------------------------------

SIMILAR SENTENCES
"""

                if similar_sentences:

                    for index, item in enumerate(
                        similar_sentences,
                        start=1
                    ):

                        report += f"""

Match {index}
Similarity: {item['score']:.2f}%

Original Sentence:
{item['original']}

Matched Sentence:
{item['matched']}

-----------------------------------
"""

                else:

                    report += """

No highly similar sentences were found.

-----------------------------------
"""

                report += f"""

-----------------------------------

FINAL ASSESSMENT

Overall Similarity: {similarity_percentage:.2f}%
Plagiarism Risk: {status}
Similar Sentences: {len(similar_sentences)}

The combined score uses lexical similarity from TF-IDF
and meaning-based similarity from Sentence Transformer
embeddings.

Generated by:
Plagiarism Detection System
"""

                st.download_button(
                    label="📥 Download TXT Report",
                    data=report,
                    file_name="plagiarism_report.txt",
                    mime="text/plain",
                    use_container_width=True
                )

                pdf_file = create_pdf_report(
                    original_name,
                    checked_name,
                    similarity_percentage,
                    status,
                    similar_sentences
                )

                st.download_button(
                    label="📄 Download PDF Report",
                    data=pdf_file,
                    file_name="plagiarism_report.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )

                docx_file = create_docx_report(
                    original_name,
                    checked_name,
                    similarity_percentage,
                    status,
                    similar_sentences
                )

                st.download_button(
                    label="📝 Download DOCX Report",
                    data=docx_file,
                    file_name="plagiarism_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )


# --------------------------------------------------
# HISTORY PAGE
# --------------------------------------------------

elif page == "📜 History":

    st.markdown(
        '<div class="main-title">📜 Analysis History</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="main-subtitle">'
        'Permanent plagiarism analysis records stored in SQLite'
        '</div>',
        unsafe_allow_html=True
    )

    history = get_history(st.session_state.user_id)

    if not history:

        st.info(
            "📭 No analysis history available yet."
        )

        st.write(
            "Run a plagiarism analysis first to create "
            "your first database record."
        )

    else:

        st.write(
            f"Total Analyses: **{len(history)}**"
        )

        csv_data = history_to_csv(history)

        st.download_button(
            label="📊 Export History as CSV",
            data=csv_data,
            file_name="plagiarism_analysis_history.csv",
            mime="text/csv",
            use_container_width=True
        )

        st.divider()

        for row in history:

            with st.expander(
                f"Analysis #{row['id']} — "
                f"{row['similarity']:.2f}% similarity"
            ):

                col1, col2, col3, col4 = st.columns(4)

                with col1:
                    st.write("📄 **Original**")
                    st.write(row["original_file"])

                with col2:
                    st.write("📄 **Checked**")
                    st.write(row["checked_file"])

                with col3:
                    st.metric(
                        "Similarity",
                        f"{row['similarity']:.2f}%"
                    )

                with col4:
                    st.metric(
                        "Risk",
                        row["risk"]
                    )

                st.write(
                    "🔎 Similar Sentences: "
                    f"**{row['similar_sentences']}**"
                )

                st.caption(
                    f"🕒 {row['created_at']}"
                )

                if st.button(
                    "🗑️ Delete This Record",
                    key=f"delete_{row['id']}",
                    use_container_width=True
                ):
                    delete_analysis(row["id"], st.session_state.user_id)
                    st.success("✅ Record deleted.")
                    st.rerun()

        st.divider()

        if st.button(
            "🗑️ Clear All History",
            use_container_width=True
        ):
            clear_history(st.session_state.user_id)
            st.success(
                "✅ All analysis history cleared."
            )
            st.rerun()


# --------------------------------------------------
# FOOTER
# --------------------------------------------------

st.markdown(
    """
    <div class="footer">
        🔍 Plagiarism Detection System
        <br>
        Built with Python, NLP, TF-IDF, Semantic Embeddings & SQLite
    </div>
    """,
    unsafe_allow_html=True
)